import streamlit as st
import pandas as pd
import os
import streamlit_authenticator as stauth
import yaml
import bcrypt
import plotly.express as px
import requests 
import json
import re
import unicodedata
import csv 
from datetime import datetime, timedelta
from io import BytesIO
from docx import Document 
from docx.shared import Pt, RGBColor

# --- CẤU HÌNH TRANG ---
st.set_page_config(page_title="BHXH Web Manager", layout="wide", initial_sidebar_state="expanded")

# --- CẤU HÌNH FILE ---
PARQUET_FILE = 'data_cache.parquet' 
EXCEL_FILE = 'aaa.xlsb' 
USER_DB_FILE = 'users.json' 
LOG_FILE = 'activity_logs.csv' 
COT_UU_TIEN = ['hoTen', 'ngaySinh', 'soBhxh', 'hanTheDen', 'soCmnd', 'soDienThoai', 'diaChiLh', 'VSS_EMAIL']

# --- HỆ THỐNG LOGGING (NHẬT KÝ) ---
def log_action(username, action, detail=""):
    """Ghi lại hoạt động của người dùng vào file CSV."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    file_exists = os.path.isfile(LOG_FILE)
    with open(LOG_FILE, mode='a', newline='', encoding='utf-8') as file:
        writer = csv.writer(file)
        if not file_exists:
            writer.writerow(['Thời gian', 'Người dùng', 'Hành động', 'Chi tiết'])
        writer.writerow([timestamp, username, action, detail])

def hien_thi_nhat_ky_he_thong(user_config):
    """Hiển thị và lọc nhật ký hoạt động của từng người dùng."""
    st.markdown("### 🕵️‍♂️ NHẬT KÝ HOẠT ĐỘNG HỆ THỐNG")
    
    if os.path.exists(LOG_FILE):
        df_log = pd.read_csv(LOG_FILE)
        df_log = df_log.sort_values(by='Thời gian', ascending=False)
        
        user_list = ['Tất cả người dùng'] + list(user_config['usernames'].keys())
        selected_user = st.selectbox("Chọn người dùng để xem nhật ký:", user_list)
        
        df_display = df_log.copy()
        
        if selected_user != 'Tất cả người dùng':
            df_display = df_log[df_log['Người dùng'] == selected_user]
            st.info(f"Đang hiển thị nhật ký của: **{selected_user}** ({len(df_display)} hoạt động).")
        else:
            st.info(f"Đang hiển thị nhật ký của: **Tất cả người dùng** ({len(df_display)} hoạt động).")
            
        if df_display.empty:
            st.warning("Không có hoạt động nào được ghi lại cho người dùng này.")
            return

        st.dataframe(df_display, use_container_width=True, height=500)
        csv = df_display.to_csv(index=False).encode('utf-8-sig')
        st.download_button(f"📥 Tải Nhật ký của {selected_user}", csv, f"nhat_ky_{selected_user}.csv", "text/csv")
        
    else:
        st.info("Chưa có nhật ký hoạt động nào.")

# --- HÀM QUẢN LÝ USER ---
def load_users():
    if not os.path.exists(USER_DB_FILE):
        hashed_pw = bcrypt.hashpw("12345".encode(), bcrypt.gensalt()).decode()
        default_data = {
            'usernames': {
                'bhxh_admin': {
                    'name': 'Admin Tổng',
                    'email': 'admin@bhxh.vn',
                    'password': hashed_pw,
                    'role': 'admin'
                }
            }
        }
        with open(USER_DB_FILE, 'w') as f: json.dump(default_data, f)
        return default_data
    try:
        with open(USER_DB_FILE, 'r') as f: return json.load(f)
    except Exception: return {}

def save_users(config):
    with open(USER_DB_FILE, 'w') as f: json.dump(config, f)

# --- GIAO DIỆN QUẢN TRỊ USER (ADMIN) ---
def hien_thi_quan_ly_user(config):
    st.markdown("### 👥 QUẢN TRỊ NGƯỜI DÙNG")
    
    tab1, tab2, tab3, tab4 = st.tabs(["➕ Thêm User", "🛠️ Reset Mật khẩu", "🔑 Đổi MK Thủ công", "❌ Xóa User"])

    # TAB 1: THÊM USER
    with tab1:
        with st.form("add_user_form"):
            c1, c2 = st.columns(2)
            new_username = c1.text_input("Tên đăng nhập (Viết liền)", placeholder="vd: nhanvien1")
            new_name = c2.text_input("Tên hiển thị", placeholder="vd: Nguyễn Văn A")
            new_password = c1.text_input("Mật khẩu khởi tạo", type="password")
            new_role = c2.selectbox("Phân quyền", ["user", "admin"], index=0)
            
            if st.form_submit_button("Lưu tài khoản"):
                if new_username and new_password and new_name:
                    if new_username in config['usernames']:
                        st.error("❌ Tên đăng nhập này đã tồn tại!")
                    else:
                        hashed_pw = bcrypt.hashpw(new_password.encode(), bcrypt.gensalt()).decode()
                        config['usernames'][new_username] = {'name': new_name, 'password': hashed_pw, 'role': new_role, 'email': ''}
                        save_users(config)
                        log_action(st.session_state["username"], "Thêm User", f"User: {new_username}")
                        st.success(f"✅ Đã tạo user: {new_username}")
                        st.rerun()
                else: st.warning("⚠️ Vui lòng điền đủ thông tin.")

    # TAB 2: RESET MẬT KHẨU VỀ MẶC ĐỊNH
    with tab2:
        st.warning("⚠️ Chức năng này sẽ đặt lại mật khẩu của user về mặc định là: **123456**")
        list_users = list(config['usernames'].keys())
        col_res_1, col_res_2 = st.columns([3, 1])
        with col_res_1: user_to_reset = st.selectbox("Chọn tài khoản cần Reset:", list_users, key="sel_reset")
        with col_res_2:
            st.write("") 
            st.write("")
            if st.button("🔄 Reset về 123456", type="primary"):
                try:
                    default_pw_hash = bcrypt.hashpw("123456".encode(), bcrypt.gensalt()).decode()
                    config['usernames'][user_to_reset]['password'] = default_pw_hash
                    save_users(config)
                    log_action(st.session_state["username"], "Reset MK", f"User: {user_to_reset}")
                    st.success(f"✅ Đã reset mật khẩu của **{user_to_reset}** thành **123456**")
                except Exception as e: st.error(f"Lỗi: {e}")

    # TAB 3: ĐỔI MẬT KHẨU THỦ CÔNG
    with tab3:
        st.info("Đổi mật khẩu thủ công sang một mật khẩu mới cụ thể.")
        list_all_users = list(config['usernames'].keys())
        
        col_change_1, col_change_2 = st.columns([3, 1])
        with col_change_1:
            user_to_change = st.selectbox("Chọn tài khoản:", list_all_users, key="sel_change")
        with col_change_2:
            st.write("") 
            st.write("") 
            if st.button("💾 Cập nhật MK"):
                if st.session_state.new_pass_change:
                    new_hash = bcrypt.hashpw(st.session_state.new_pass_change.encode(), bcrypt.gensalt()).decode()
                    config['usernames'][user_to_change]['password'] = new_hash
                    save_users(config)
                    log_action(st.session_state["username"], "Đổi MK thủ công", f"User: {user_to_change}")
                    st.success(f"✅ Đã đổi mật khẩu cho: {user_to_change}")
                else: st.error("Chưa nhập mật khẩu.")
        st.text_input("Nhập mật khẩu mới:", type="password", key="new_pass_change") # Input bên ngoài form

    # TAB 4: XÓA USER
    with tab4:
        st.error("⚠️ Hành động xóa không thể hoàn tác.")
        current_user = st.session_state["username"]
        list_users_to_delete = [u for u in config['usernames'].keys() if u != current_user]
        
        if list_users_to_delete:
            user_to_delete = st.selectbox("Chọn tài khoản cần xóa:", list_users_to_delete, key="sel_del")
            if st.button("🗑️ Xác nhận xóa", type="primary"):
                try:
                    del config['usernames'][user_to_delete]
                    save_users(config)
                    log_action(st.session_state["username"], "Xóa User", f"User: {user_to_delete}")
                    st.success(f"✅ Đã xóa tài khoản: {user_to_delete}")
                    st.rerun()
                except Exception as e: st.error(f"Lỗi: {e}")
        else: st.info("Không có tài khoản nào khác để xóa.")

    # HIỂN THỊ DANH SÁCH
    st.divider()
    st.subheader("Danh sách tài khoản hiện có")
    user_list = []
    for u, data in config['usernames'].items():
        user_list.append({"Tên đăng nhập": u, "Tên hiển thị": data['name'], "Quyền hạn": data.get('role', 'user'), "Trạng thái mật khẩu": "Đã mã hóa (Ẩn)"})
    st.dataframe(pd.DataFrame(user_list), use_container_width=True)

# --- GIAO DIỆN QUẢN TRỊ CHUNG (ADMIN) ---
def hien_thi_quan_tri_admin(config):
    st.markdown("### ⚙️ TRUNG TÂM QUẢN TRỊ")
    
    st.markdown("#### Quản lý Người dùng")
    col_u1, col_u2 = st.columns(2)
    col_u1.button("👥 QUẢN LÝ USER", on_click=set_state, args=('admin_user',))
    col_u2.button("📝 NHẬT KÝ", on_click=set_state, args=('admin_log',)) 
    
    st.markdown("#### Quản lý Dữ liệu")
    st.button("⚙️ CẬP NHẬT DATA", on_click=set_state, args=('admin_data',))
    
    st.divider()

    if st.session_state.get('admin_user'): hien_thi_quan_ly_user(config)
    elif st.session_state.get('admin_data'): hien_thi_quan_tri_data()
    elif st.session_state.get('admin_log'): hien_thi_nhat_ky_he_thong(config)
    else: st.info("Chọn một chức năng quản trị bên trên.")

# --- HÀM HỖ TRỢ & NẠP DỮ LIỆU (GIỮ NGUYÊN) ---
def xoa_dau_tieng_viet(text):
    if not isinstance(text, str): return str(text)
    text = unicodedata.normalize('NFD', text)
    text = re.sub(r'[\u0300-\u036f]', '', text)
    text = text.lower().strip()
    text = re.sub(r'\s+', ' ', text)
    return text

def set_state(name):
    for key in ['search', 'loc', 'han', 'bieu', 'chuan', 'ai', 'admin_data', 'admin_user', 'admin_log', 'admin_panel']:
        st.session_state[key] = False
    st.session_state[name] = True

def hien_thi_quan_tri_data():
    st.markdown("### ⚙️ CẬP NHẬT DỮ LIỆU HỆ THỐNG")
    uploaded_file = st.file_uploader("📂 Chọn file Excel dữ liệu (.xlsb)", type=['xlsb'])
    if uploaded_file is not None:
        if st.button("🚀 CẬP NHẬT DỮ LIỆU"):
            try:
                with st.spinner("Đang xử lý..."):
                    with open(EXCEL_FILE, "wb") as f: f.write(uploaded_file.getbuffer())
                    st.cache_data.clear()
                    if os.path.exists(PARQUET_FILE): os.remove(PARQUET_FILE)
                    nap_du_lieu_toi_uu()
                    log_action(st.session_state["username"], "Cập nhật Data", f"File: {uploaded_file.name}")
                st.success("✅ Cập nhật thành công!")
            except Exception as e: st.error(f"Có lỗi xảy ra: {e}")

def tao_phieu_word(row):
    doc = Document()
    heading = doc.add_heading('PHIẾU THÔNG TIN BHXH', 0); heading.alignment = 1
    doc.add_paragraph(f'Ngày xuất phiếu: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    doc.add_paragraph('--------------------------------------------------')
    p = doc.add_paragraph(); run = p.add_run(f"HỌ VÀ TÊN: {row.get('hoTen', '').upper()}")
    run.bold = True; run.font.size = Pt(14); run.font.color.rgb = RGBColor(0, 51, 102)
    table = doc.add_table(rows=1, cols=2); table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells; hdr_cells[0].text = 'THÔNG TIN'; hdr_cells[1].text = 'CHI TIẾT'
    for cot in COT_UU_TIEN:
        row_cells = table.add_row().cells; row_cells[0].text = cot
        val = row.get(cot, ''); row_cells[1].text = str(val) if pd.notna(val) else ""
    doc.add_paragraph('\n'); doc.add_paragraph('Người trích xuất: Admin BHXH').alignment = 2
    bio = BytesIO(); doc.save(bio); return bio

def tao_file_excel(df_input):
    output = BytesIO(); writer = pd.ExcelWriter(output, engine='xlsxwriter')
    df_input.to_excel(writer, index=False, sheet_name='DanhSach')
    writer.close(); return output

@st.cache_data(ttl=3600)
def nap_du_lieu_toi_uu():
    if os.path.exists(PARQUET_FILE):
        try:
            df = pd.read_parquet(PARQUET_FILE)
            cols_to_str = ['soBhxh', 'soCmnd', 'soDienThoai', 'ngaySinh', 'hanTheDen']
            for col in cols_to_str:
                if col in df.columns: df[col] = df[col].astype(str)
            return df
        except Exception: pass
    if not os.path.exists(EXCEL_FILE): return pd.DataFrame()
    try:
        with st.spinner('⚙️ Đang tối ưu hóa dữ liệu...'):
            df = pd.read_excel(EXCEL_FILE, dtype=str, engine='pyxlsb')
            df.columns = df.columns.str.strip()
            df.to_parquet(PARQUET_FILE)
        return df
    except Exception as e: return pd.DataFrame()

def hien_thi_uu_tien(df_ket_qua):
    if df_ket_qua.empty: st.warning("😞 Không tìm thấy kết quả phù hợp."); return
    st.success(f"✅ Tìm thấy {len(df_ket_qua)} hồ sơ!"); excel_data = tao_file_excel(df_ket_qua)
    st.download_button(label="📥 Tải danh sách (Excel)", data=excel_data.getvalue(), file_name=f"danh_sach.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    if len(df_ket_qua) > 50: st.caption(f"⚠️ Đang hiển thị 50/{len(df_ket_qua)} kết quả đầu tiên."); df_ket_qua = df_ket_qua.head(50)
    for i in range(min(len(df_ket_qua), 50)):
        row = df_ket_qua.iloc[i]; tieu_de = f"👤 {row.get('hoTen', 'Na')} - {row.get('soBhxh', '')}"
        with st.expander(tieu_de, expanded=False):
            c1, c2 = st.columns([3, 1]); 
            with c1:
                col_a, col_b = st.columns(2);
                for idx, cot in enumerate(COT_UU_TIEN):
                    val = "(Trống)"; 
                    for c_ex in df_ket_qua.columns:
                         if cot.lower() == c_ex.lower():
                             v = row[c_ex]; val = str(v) if pd.notna(v) and str(v).strip() != "" and str(v).lower() != "nan" else "(Trống)"; break
                    if idx % 2 == 0: col_a.markdown(f"**🔹 {cot}:** {val}")
                    else: col_b.markdown(f"**🔹 {cot}:** {val}")
            with c2: w_data = tao_phieu_word(row); st.download_button(label="📄 In Phiếu", data=w_data.getvalue(), file_name=f"Phieu_{row.get('soBhxh', 'hs')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"btn_word_{i}")
            st.dataframe(row.to_frame().T, hide_index=True)

def hien_thi_loc_loi(df, ten_cot):
    log_action(st.session_state["username"], "Lọc Lỗi", f"Cột: {ten_cot}"); col_chuan = df[ten_cot].astype(str).str.strip().str.lower(); rong = ['nan', 'none', 'null', '', '0']; df_loc = df[col_chuan.isin(rong)]
    if not df_loc.empty:
        st.warning(f"⚠️ {len(df_loc)} hồ sơ thiếu '{ten_cot}'."); excel_data = tao_file_excel(df_loc)
        st.download_button(label="📥 Tải danh sách lỗi", data=excel_data.getvalue(), file_name=f"loi_{ten_cot}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"); st.dataframe(df_loc.head(1000))
    else: st.success(f"Tuyệt vời! Cột '{ten_cot}' đủ dữ liệu.")

def hien_thi_kiem_tra_han(df, ten_cot_ngay):
    log_action(st.session_state["username"], "Kiểm tra hạn", ten_cot_ngay); df_temp = df[[ten_cot_ngay, 'hoTen', 'soBhxh']].copy()
    try:
        df_temp[ten_cot_ngay] = pd.to_datetime(df_temp[ten_cot_ngay], dayfirst=True, errors='coerce'); df_co = df_temp.dropna(subset=[ten_cot_ngay])
        hom_nay = datetime.now(); sau_30 = hom_nay + timedelta(days=30); ds_het = df_co[df_co[ten_cot_ngay] < hom_nay].copy()
        ds_sap = df_co[(df_co[ten_cot_ngay] >= hom_nay) & (df_co[ten_cot_ngay] <= sau_30)].copy()
        if not ds_het.empty: ds_het[ten_cot_ngay] = ds_het[ten_cot_ngay].dt.strftime('%d/%m/%Y')
        if not ds_sap.empty: ds_sap[ten_cot_ngay] = ds_sap[ten_cot_ngay].dt.strftime('%d/%m/%Y')
        c1, c2 = st.columns(2); c1.metric("🔴 ĐÃ HẾT HẠN", f"{len(ds_het)}"); c2.metric("⚠️ SẮP HẾT HẠN", f"{len(ds_sap)}")
        if not ds_het.empty:
            st.subheader("🔴 Danh sách Hết Hạn"); excel_het = tao_file_excel(ds_het); st.download_button("📥 Tải Hết Hạn", excel_het.getvalue(), "het_han.xlsx"); st.dataframe(ds_het.head(500), hide_index=True)
        if not ds_sap.empty:
            st.subheader("⚠️ Danh sách Sắp Hết"); excel_sap = tao_file_excel(ds_sap); st.download_button("📥 Tải Sắp Hết", excel_sap.getvalue(), "sap_het.xlsx"); st.dataframe(ds_sap.head(500), hide_index=True)
    except Exception as e: st.error(f"Lỗi ngày tháng: {e}")

def hien_thi_bieu_do_tuong_tac(df, ten_cot):
    log_action(st.session_state["username"], "Xem Biểu Đồ", ten_cot); st.markdown(f"### 📊 BIỂU ĐỒ TƯƠNG TÁC: {ten_cot}")
    thong_ke = df[ten_cot].value_counts().reset_index(); thong_ke.columns = ['Phân loại', 'Số lượng'] 
    fig = px.bar(thong_ke, x='Phân loại', y='Số lượng', text='Số lượng', color='Phân loại')
    event = st.plotly_chart(fig, use_container_width=True, on_select="rerun")
    if event and event['selection']['points']:
        gia_tri_chon = event['selection']['points'][0]['x']; st.divider(); st.info(f"🔍 Bạn vừa chọn: **{gia_tri_chon}**.")
        log_action(st.session_state["username"], "Click Biểu Đồ", f"Xem chi tiết: {gia_tri_chon}"); hien_thi_uu_tien(df[df[ten_cot] == gia_tri_chon])

def hien_thi_chatbot_thong_minh(df):
    log_action(st.session_state["username"], "Xem Chatbot", ""); st.markdown("### 🤖 TRỢ LÝ ẢO (Tìm Kiếm Linh Hoạt)")
    if "messages" not in st.session_state: st.session_state.messages = []
    for msg in st.session_state.messages: 
        # FIX: Dòng lỗi (line 332) đã được sửa lại cú pháp (with/as) chuẩn
        with st.chat_message(msg["role"]): st.markdown(msg["content"])
        
    if prompt := st.chat_input("Nhập yêu cầu..."):
        st.session_state.messages.append({"role": "user", "content": prompt}); with st.chat_message("user"): st.markdown(prompt); log_action(st.session_state["username"], "Chat AI", prompt)
        with st.chat_message("assistant"):
            df_res = df.copy(); df_res['hoTen_khongdau'] = df_res['hoTen'].apply(lambda x: xoa_dau_tieng_viet(str(x))); filters = [] 
            try:
                date_m = re.search(r'\d{1,2}[/-]\d{1,2}[/-]\d{4}', prompt);
                if date_m:
                    ngay_raw = date_m.group().replace('-', '/');
                    try:
                        nd = pd.to_datetime(ngay_raw, dayfirst=True).strftime('%d/%m/%Y'); mask_date = df_res['ngaySinh'].astype(str).str.contains(nd); df_res = df_res[mask_date]; filters.append(f"Ngày sinh: **{nd}**");
                    except: pass
                nums = re.findall(r'\b\d{5,}\b', prompt);
                for n in nums:
                    if date_m and n in date_m.group(): continue; mask_so = (df_res['soBhxh'].astype(str).str.contains(n)) | (df_res['soCmnd'].astype(str).str.contains(n)); df_res = df_res[mask_so]; filters.append(f"Mã: {n}")
                tu_rac = ["tim", "loc", "cho", "toi", "nguoi", "co", "ngay", "sinh", "ten", "la", "o", "que"];
                p_clean = xoa_dau_tieng_viet(prompt); for w in tu_rac: p_clean = re.sub(r'\b' + w + r'\b', '', p_clean);
                p_clean = re.sub(r'\b(bieu do|thong ke|han|het han)\b', '', p_clean); ten = re.sub(r'\s+', ' ', p_clean).strip();
                if len(ten) > 1:
                    df_res = df_res[df_res['hoTen_khongdau'].str.contains(ten)]; filters.append(f"Tên: {ten}")
                if "bieu do" in xoa_dau_tieng_viet(prompt):
                    cot_ve = 'gioiTinh'; 
                    if "tinh" in xoa_dau_tieng_viet(prompt): cot_ve = 'maTinh';
                    st.write(f"📈 Biểu đồ: {cot_ve}"); hien_thi_bieu_do_tuong_tac(df, cot_ve)
                elif "han" in xoa_dau_tieng_viet(prompt) and "het" in xoa_dau_tieng_viet(prompt):
                    st.write("⏳ Kiểm tra hạn BHYT..."); hien_thi_kiem_tra_han(df, 'hanTheDen')
                elif filters:
                    st.write(f"🔍 Điều kiện: {' + '.join(filters)}")
                    if not df_res.empty:
                        if 'hoTen_khongdau' in df_res.columns: df_res = df_res.drop(columns=['hoTen_khongdau'])
                        hien_thi_uu_tien(df_res)
                    else: st.warning("Không tìm thấy ai.")
                else: st.info("🤖 Hãy nhập tên hoặc ngày sinh để tìm kiếm.")
            except Exception as e: st.error(f"Lỗi: {e}")

# --- MAIN ---
def main():
    user_config = load_users()
    authenticator = stauth.Authenticate(user_config, 'bhxh_cookie', 'key_bi_mat_rat_dai_va_kho_doan_123', 30)
    
    authenticator.login(location='main') 

    if st.session_state["authentication_status"]:
        if 'logged_in' not in st.session_state:
            log_action(st.session_state["username"], "Đăng nhập", "Thành công")
            st.session_state['logged_in'] = True

        username = st.session_state["username"]
        user_role = user_config['usernames'][username].get('role', 'user')
        user_name_display = user_config['usernames'][username]['name']

        with st.sidebar:
            st.write(f'Xin chào, **{user_name_display}**! 👋')
            if user_role == 'admin': st.caption("👑 Quản trị viên")
            else: st.caption("👤 Người dùng")
            
            authenticator.logout('Đăng xuất', 'sidebar')
            st.markdown("---")
        
        st.title("🌐 HỆ THỐNG QUẢN LÝ BHXH")
        df = nap_du_lieu_toi_uu() 
        
        if df.empty:
            st.warning("⚠️ Chưa có dữ liệu.")
            if user_role == 'admin':
                st.sidebar.button("⚙️ QUẢN TRỊ DATA", on_click=set_state, args=('admin_data',))
                if st.session_state.get('admin_data'): hien_thi_quan_tri_data()
            return

        st.sidebar.header("CHỨC NĂNG")
        cols = df.columns.tolist()
        idx_sobhxh = cols.index('soBhxh') if 'soBhxh' in cols else 0
        ten_cot = st.sidebar.selectbox("Cột xử lý:", options=cols, index=idx_sobhxh)
        tim_kiem = st.sidebar.text_input("Tìm kiếm nhanh:", placeholder="Nhập tên...")

        st.sidebar.markdown("---")
        c1, c2 = st.sidebar.columns(2)
        c1.button("🔍 TRA CỨU", on_click=set_state, args=('search',))
        c2.button("🧹 LỌC LỖI", on_click=set_state, args=('loc',))
        c3, c4 = st.sidebar.columns(2)
        c3.button("⏳ HẠN BHYT", on_click=set_state, args=('han',))
        c4.button("📊 BIỂU ĐỒ", on_click=set_state, args=('bieu',))
        st.sidebar.markdown("---")
        st.sidebar.button("🤖 TRỢ LÝ ẢO", on_click=set_state, args=('ai',))
        
        if user_role == 'admin':
            st.sidebar.markdown("---")
            st.sidebar.caption("QUẢN TRỊ HỆ THỐNG")
            st.sidebar.button("📝 NHẬT KÝ", on_click=set_state, args=('admin_log',)) 
            st.sidebar.button("⚙️ CẬP NHẬT DATA", on_click=set_state, args=('admin_data',))
            st.sidebar.button("👥 QUẢN LÝ USER", on_click=set_state, args=('admin_user',))

        st.markdown("---")
        for key in ['search', 'loc', 'han', 'bieu', 'ai', 'admin_data', 'admin_user', 'admin_log']:
            if key not in st.session_state: st.session_state[key] = False

        if st.session_state.get('loc'): hien_thi_loc_loi(df, ten_cot)
        elif st.session_state.get('han'): hien_thi_kiem_tra_han(df, ten_cot)
        elif st.session_state.get('bieu'): hien_thi_bieu_do_tuong_tac(df, ten_cot)
        elif st.session_state.get('ai'): hien_thi_chatbot_thong_minh(df)
        elif st.session_state.get('admin_data') and user_role == 'admin': hien_thi_quan_tri_data()
        elif st.session_state.get('admin_user') and user_role == 'admin': hien_thi_quan_ly_user(user_config)
        elif st.session_state.get('admin_log') and user_role == 'admin': hien_thi_nhat_ky_he_thong(user_config)
        
        elif tim_kiem:
            log_action(username, "Tìm kiếm nhanh", f"Từ khóa: {tim_kiem} (Cột: {ten_cot})")
            mask = df[ten_cot].astype(str).str.contains(tim_kiem, case=False, na=False)
            hien_thi_uu_tien(df[mask])
        else:
            st.info("👈 Chọn chức năng bên trái.")
            st.caption("Dữ liệu mẫu:")
            st.dataframe(df.head(10))

    elif st.session_state["authentication_status"] is False: st.error('Sai mật khẩu.')
    elif st.session_state["authentication_status"] is None: st.warning('Vui lòng đăng nhập.')

if __name__ == "__main__":
    main()